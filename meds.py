import tkinter as tk
from tkinter import ttk, messagebox
import tkcalendar  # Add this import for the calendar widget
from db_helper import DatabaseHelper
from medicine_select import MedicineSelector
from medical_certificate import MedicalCertificateWindow  # Import the new class
import sqlite3
from datetime import datetime, date
from tkcalendar import Calendar  # Change from DateEntry to Calendar
from medication_management import MedicationManagementWindow
import os
import tempfile
import subprocess
from PIL import Image, ImageDraw, ImageFont
from tkinter import filedialog
from lab_charts import LabChartsWindow

# Replace the AutocompleteCombobox class with this updated version
class AutocompleteCombobox(ttk.Combobox):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._is_selecting = False
        
        # Override the standard bindings to disable automatic dropdown
        self.bind("<KeyRelease>", self._on_key_release)
        self.bind("<<ComboboxSelected>>", self._on_selection)
        self.bind("<Down>", self._on_down)
        
        # Disable the built-in Combobox postcommand that shows the dropdown
        self._original_postcommand = self.cget('postcommand')
        self.config(postcommand=self._prevent_auto_post)
        
    def _prevent_auto_post(self):
        """Override default dropdown behavior to prevent it from showing automatically"""
        # Call original postcommand to update values, but don't show dropdown
        if callable(self._original_postcommand):
            self._original_postcommand()
    
    def _on_key_release(self, event):
        """Handle key release to update matches but not show dropdown"""
        # Skip processing for navigational keys
        if event.keysym in ('Up', 'Down', 'Left', 'Right', 'Return', 'Escape'):
            return
        
        # Update matches without showing dropdown
        self.after(1, lambda: check_name_match(event))
    
    def _on_selection(self, event):
        """Handle selection when dropdown is explicitly shown"""
        self._is_selecting = True
        # Call the original selection handler
        on_name_select(event)
        self._is_selecting = False
        return "break"  # Prevent default behavior
    
    def _on_down(self, event):
        """Custom down arrow handler to allow explicit dropdown show"""
        # Only show dropdown when Down arrow is pressed
        if self.cget('values'):
            self.event_generate('<Down>')
            return "break"  # Handle it ourselves
        return  # Let default handler work

# Update the check_name_match function to populate values but NOT show dropdown
def check_name_match(event=None):
    """Filter dropdown list based on typed text but don't show the dropdown"""
    current_text = entry_name.get().strip()
    
    # When text is empty, set values but don't show dropdown
    if not current_text:
        sorted_names = sorted(list(patient_dict.keys()))[:20]  # Limit to 20 names
        entry_name['values'] = sorted_names
        return
    
    # Filter patient names based on typed text
    matching_names = []
    
    # First priority: names that start with the typed text
    for name in patient_dict.keys():
        if name.lower().startswith(current_text.lower()):
            matching_names.append(name)
    
    # Second priority: names containing the typed text (if few starting matches)
    if len(matching_names) < 5:
        for name in patient_dict.keys():
            if not name.lower().startswith(current_text.lower()) and current_text.lower() in name.lower():
                matching_names.append(name)
    
    # Update values but DO NOT show dropdown
    if matching_names:
        matching_names.sort()
        entry_name['values'] = matching_names
    else:
        entry_name['values'] = ["No matches found"]

# Add a new function to allow showing dropdown on demand
def show_patient_dropdown(event=None):
    """Function to explicitly show the patient dropdown when requested"""
    if entry_name['values']:
        entry_name.event_generate('<Down>')

# Main Application Window
root = tk.Tk()
root.title("Clinic Management System")
root.geometry("1200x700")  # Increased width to accommodate sidebar
root.configure(bg="#e6f7ff")  # Light blue background


# Color scheme
PRIMARY_COLOR = "#3498db"  # Blue
SECONDARY_COLOR = "#e6f7ff"  # Light blue
ACCENT_COLOR = "#2ecc71"  # Green
WARNING_COLOR = "#e74c3c"  # Red
TEXT_COLOR = "#2c3e50"  # Dark blue/grey
BUTTON_TEXT_COLOR = "white"
SIDEBAR_COLOR = "#2c3e50"  # Dark blue/grey for the sidebar

# Custom style
style = ttk.Style()
style.configure("TEntry", padding=5)
style.configure("TButton", padding=5)
style.configure("Treeview", rowheight=25)
style.configure("Treeview.Heading", font=("Arial", 10, "bold"))

# ----------------- Navigation Sidebar ----------------- #
sidebar = tk.Frame(root, bg=SIDEBAR_COLOR, width=180)
sidebar.pack(side=tk.LEFT, fill=tk.Y)
sidebar.pack_propagate(False)  # Prevents the sidebar from shrinking

# Logo or system title
tk.Label(sidebar, text="CLINIC\nMANAGEMENT", bg=SIDEBAR_COLOR, fg="white", 
         font=("Arial", 12, "bold"), pady=20).pack(fill=tk.X)

# Separator
ttk.Separator(sidebar).pack(fill=tk.X, padx=10)

# Function to create sidebar buttons
def create_sidebar_button(text, command=None, color=PRIMARY_COLOR):
    btn = tk.Button(sidebar, text=text, bg=color, fg=BUTTON_TEXT_COLOR, 
                   padx=5, pady=10, bd=0, width=15, command=command,
                   font=("Arial", 10, "bold"))
    btn.pack(pady=5, padx=10)
    return btn

# Content frame to hold all other widgets
content_frame = tk.Frame(root, bg=SECONDARY_COLOR)
content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Button function definitions
def save_record():
    try:
        db = DatabaseHelper()
        patient_name = entry_name.get()
        
        if not patient_name:
            messagebox.showwarning("Warning", "Please enter a patient name.")
            return
            
        # Check if patient already exists
        existing_patient = db.get_patient_by_name(patient_name)
        patient_id = None
        is_new_patient = False
        
        if existing_patient:
            patient_id = existing_patient[0]  # Get existing patient's ID
            
            # Update patient's basic information
            patient_data = (
                entry_name.get(),
                entry_address.get(),
                selected_date.get(),
                entry_phone.get(),
                status_var.get(),
                gender_var.get(),
                patient_id  # Last parameter for WHERE clause in update query
            )
            db.update_patient(patient_data)
        else:
            # Save new patient data
            is_new_patient = True
            patient_data = (
                patient_name,
                entry_address.get(),
                selected_date.get(),
                entry_phone.get(),
                status_var.get(),
                gender_var.get()
            )
            patient_id = db.add_patient(patient_data)
        
        # Check for existing checkup for today
        current_date = datetime.now().strftime('%Y-%m-%d')
        existing_checkup = db.get_checkup_by_date(patient_id, current_date)
        
        if existing_checkup:
            # Update existing checkup for today
            checkup_data = (
                text_remarks.get("1.0", tk.END),  # findings
                existing_checkup[2],  # Keep existing lab_ids
                entry_bp.get(),  # blood_pressure
                existing_checkup[0]  # checkup_id
            )
            db.update_checkup(checkup_data)
            
            # Delete existing prescriptions for updating
            db.delete_prescriptions_for_checkup(patient_id, current_date)
        else:
            # Create a new checkup record
            checkup_data = (
                patient_id,
                text_remarks.get("1.0", tk.END),  # findings
                "",          # lab_ids
                current_date,  # dateOfVisit
                current_date   # last_checkup_date
            )
            db.add_checkup(checkup_data)
        
        # Save prescriptions with current date
        for item in tree_med.get_children():
            values = tree_med.item(item)['values']
            prescription_data = (
                patient_id,
                values[1],  # generic
                values[0],  # brand
                values[2],  # quantity
                values[3],  # administration
                current_date  # last_checkup_date
            )
            db.add_prescription(prescription_data)
        
        # Display success message
        if is_new_patient:
            messagebox.showinfo("Success", "New patient record saved successfully!")
        else:
            messagebox.showinfo("Success", "Patient record updated successfully!")
        
        # Refresh the patient list and keep the current patient selected
        refresh_patient_list(patient_name)
        load_checkup_history(patient_id)
        
        # Show a visual confirmation that data is refreshed
        status_label = tk.Label(frame_patient, text="✓ Patient data saved", 
                             bg=ACCENT_COLOR, fg="white", font=("Arial", 10))
        status_label.grid(row=5, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        
        # Remove the status label after 2 seconds
        root.after(2000, status_label.destroy)
        
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def load_today_queue():
    global queue_counter
    try:
        db = DatabaseHelper()
        
        # Ensure database tables are created
        db.create_tables()
        
        queue_entries = db.get_todays_queue()
        
        # Clear existing entries in the treeview
        for item in tree_queue.get_children():
            tree_queue.delete(item)
            
        # Find highest queue number to set counter
        highest_queue = 0
        
        # Add queue entries to the treeview
        for entry in queue_entries:
            queue_id, queue_num, name, queue_time = entry
            tree_queue.insert("", "end", values=(queue_num, name, queue_time), tags=(str(queue_id),))
            highest_queue = max(highest_queue, int(queue_num) if isinstance(queue_num, (int, str)) else 0)
        
        # Set queue counter to continue from highest number
        queue_counter = highest_queue
        
    except Exception as e:
        print(f"Error loading queue: {e}")
      

  
def update_record():
    try:
        db = DatabaseHelper()
        patient_name = entry_name.get()
        
        # Check if the patient exists
        if not patient_name or patient_name not in patient_dict:
            messagebox.showwarning("Error", "Please select a valid patient first.")
            return
        
        patient_id = patient_dict[patient_name]
        
        # Get the selected date from the dropdown or use current date
        selected_checkup_date = checkup_history_var.get()
        if selected_checkup_date == "No previous checkups":
            # No checkup selected, use today's date
            checkup_date = datetime.now().strftime('%Y-%m-%d')
            is_new_checkup = True
        else:
            # Use the selected date
            checkup_date = selected_checkup_date
            is_new_checkup = False
        
        # For existing checkups, get the checkup ID
        if not is_new_checkup:
            existing_checkup = next((c for c in current_checkups if c[4] == checkup_date), None)
            if existing_checkup:
                checkup_id = existing_checkup[0]
                
                # Update the existing checkup record
                checkup_data = (
                    text_remarks.get("1.0", tk.END),  # findings
                    existing_checkup[2],  # Keep existing lab_ids
                    entry_bp.get(),      # blood_pressure
                    checkup_id           # checkup_id
                )
                db.update_checkup(checkup_data)
                
                # Delete existing prescriptions for this date
                db.delete_prescriptions_for_checkup(patient_id, checkup_date)
                
                # Add new prescriptions
                for item in tree_med.get_children():
                    values = tree_med.item(item)['values']
                    prescription_data = (
                        patient_id,
                        values[1],  # generic
                        values[0],  # brand
                        values[2],  # quantity
                        values[3],  # administration
                        checkup_date  # last_checkup_date
                    )
                    db.add_prescription(prescription_data)
            else:
                messagebox.showerror("Error", "Selected checkup record not found.")
                return
        else:
            # Create a new checkup for today
            existing_today_checkup = db.get_checkup_by_date(patient_id, checkup_date)
            
            if existing_today_checkup:
                # If there's already a checkup for today, confirm before overwriting
                response = messagebox.askyesno(
                    "Checkup Exists", 
                    f"A checkup record already exists for today. Do you want to update it?",
                    icon='warning'
                )
                if not response:
                    return
                
                # Update existing checkup for today
                checkup_data = (
                    text_remarks.get("1.0", tk.END),  # findings
                    existing_today_checkup[2],  # Keep existing lab_ids
                    entry_bp.get(),  # blood_pressure
                    existing_today_checkup[0]  # checkup_id
                )
                db.update_checkup(checkup_data)
                
                # Delete and recreate prescriptions
                db.delete_prescriptions_for_checkup(patient_id, checkup_date)
            else:
                # Create a new checkup record
                checkup_data = (
                    patient_id,
                    text_remarks.get("1.0", tk.END),  # Use remarks as findings
                    "",          # lab_ids
                    checkup_date,  # dateOfVisit
                    checkup_date  # last_checkup_date
                )
                db.add_checkup(checkup_data)
            
            # Add prescriptions with updated format
            for item in tree_med.get_children():
                values = tree_med.item(item)['values']
                prescription_data = (
                    patient_id,
                    values[1],  # generic
                    values[0],  # brand
                    values[2],  # quantity
                    values[3],  # administration
                    checkup_date  # last_checkup_date
                )
                db.add_prescription(prescription_data)
        
        # Display a success message
        messagebox.showinfo("Success", "Record updated successfully!")
        
        # Refresh the patient data
        refresh_patient_list(patient_name)
        load_checkup_history(patient_id)
        
        # Show a visual confirmation that data is refreshed
        status_label = tk.Label(frame_patient, text="✓ Patient data updated", 
                              bg=ACCENT_COLOR, fg="white", font=("Arial", 10))
        status_label.grid(row=5, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        
        # Remove the status label after 3 seconds
        root.after(2000, status_label.destroy)
        
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred during update: {str(e)}")

def clear_form(show_message=True):
    # Clear all entry fields and text widgets
    entry_name.delete(0, tk.END)
    entry_address.delete(0, tk.END)
    entry_age.delete(0, tk.END)
    entry_phone.delete(0, tk.END)
    entry_bp.delete(0, tk.END)
    text_remarks.delete("1.0", tk.END)
    # Clear medications treeview
    for item in tree_med.get_children():
        tree_med.delete(item)
    # Reset comboboxes and radiobuttons
    status_var.set("")
    gender_var.set("")
    selected_date.set(datetime.now().strftime('%Y-%m-%d'))
        # Show message only if show_message is True
    if show_message:
        messagebox.showinfo("Clear", "Form cleared successfully!")

def delete_record():
    # Check if a patient is selected
    patient_name = entry_name.get()
    if not patient_name:
        messagebox.showwarning("Selection Error", "Please select a patient to delete.")
        return
    
    # Check if patient exists in the database
    if patient_name not in patient_dict:
        messagebox.showwarning("Patient Not Found", f"Patient '{patient_name}' not found in the database.")
        return
    
    # Get patient ID
    patient_id = patient_dict[patient_name]
    
    # Confirm deletion with the user
    response = messagebox.askyesno(
        "Confirm Deletion", 
        f"Are you sure you want to delete ALL records for patient '{patient_name}'?\n\nThis action cannot be undone and will delete:\n- Patient information\n- All checkup records\n- All prescription history",
        icon='warning'
    )
    
    if response:
        try:
            db = DatabaseHelper()
            db.delete_patient(patient_id)
            messagebox.showinfo("Success", f"Patient '{patient_name}' and all related records have been deleted.")
            
            # Clear the form
            clear_form()
            
            # Refresh the patient list
            refresh_patient_list()
            
            # Show a visual confirmation
            status_label = tk.Label(frame_patient, text="✓ Patient deleted and list refreshed", 
                                  bg=WARNING_COLOR, fg="white", font=("Arial", 10))
            status_label.grid(row=5, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
            
            # Remove the status label after 3 seconds
            root.after(3000, status_label.destroy)
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while deleting the patient: {str(e)}")

def open_med_cert():
    # Check if a patient is selected
    if not entry_name.get():
        messagebox.showwarning("Warning", "Please select a patient first.")
        return
    
    # Gather patient data from the current form
    patient_data = {
        "name": entry_name.get(),
        "age": entry_age.get(),
        "address": entry_address.get(),
        "findings": text_remarks.get("1.0", tk.END).strip(),  # Get the content from diagnosis field
        "remarks": ""  # Leave remarks empty since we're not using it
    }
    
    # Create and open the medical certificate window
    med_cert_window = MedicalCertificateWindow(root, patient_data)


def open_print_dialog():
    """Open a print dialog window to print prescription or findings"""
    # Check if a patient is selected
    if not entry_name.get():
        messagebox.showwarning("Warning", "Please select a patient first.")
        return
    
    # Create a print dialog window
    print_dialog = tk.Toplevel(root)
    print_dialog.title("Print Document")
    print_dialog.geometry("800x600")
    print_dialog.configure(bg=SECONDARY_COLOR)
    print_dialog.resizable(False, False)
    
    # Create the content frame
    content_frame = tk.Frame(print_dialog, bg=SECONDARY_COLOR, padx=20, pady=15)
    content_frame.pack(fill=tk.BOTH, expand=True)
    
    # Print type selection
    tk.Label(content_frame, text="Print Type:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 10, "bold")).grid(row=0, column=0, padx=5, pady=10, sticky="w")
    print_type_var = tk.StringVar(value="Prescription")
    print_type_combo = ttk.Combobox(content_frame, textvariable=print_type_var, width=20, state="readonly", values=["Prescription", "Findings"])
    print_type_combo.grid(row=0, column=1, padx=5, pady=10, sticky="w")
    
    # Preview area
    preview_frame = tk.LabelFrame(content_frame, text="Print Preview", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 10, "bold"), padx=10, pady=10)
    preview_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=10, sticky="nsew")
    content_frame.grid_rowconfigure(1, weight=1)
    content_frame.grid_columnconfigure(0, weight=1)
    content_frame.grid_columnconfigure(1, weight=1)
    
    preview_text = tk.Text(preview_frame, wrap=tk.WORD, width=40, height=10, bg="white", fg=TEXT_COLOR)
    preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    
    # Function to update preview based on selected print type
    def update_preview(*args):
        preview_text.delete(1.0, tk.END)
        selected_type = print_type_var.get()
        
        # Patient name at the top
        preview_text.insert(tk.END, f"Patient: {entry_name.get()}\n\n")
        
        if selected_type == "Prescription":
            # Format prescription data
            prescription_text = get_formatted_prescription(entry_name.get())
            preview_text.insert(tk.END, prescription_text)
        else:  # Findings
            # Get the findings/remarks
            findings_text = text_remarks.get("1.0", tk.END).strip()
            preview_text.insert(tk.END, findings_text)
    
    # Bind the update_preview function to combobox selection changes
    print_type_combo.bind("<<ComboboxSelected>>", update_preview)
    
    # Button frame
    button_frame = tk.Frame(print_dialog, bg=SECONDARY_COLOR, padx=20, pady=15)
    button_frame.pack(fill=tk.X)
    
    # Define function to print as Word document
    def print_document_as_word():
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import tempfile
            import os

            # Get content based on print type
            selected_type = print_type_var.get()
            fd, docx_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            doc = Document()
            
            # Get patient information
            patient_name = entry_name.get().upper()
            patient_age = entry_age.get()
            patient_address = entry_address.get()
            current_date = datetime.now().strftime("%B %d, %Y")
            
            # Get patient gender - first character only (M or F)
            patient_gender = ""
            if gender_var.get():
                patient_gender = gender_var.get()[0]  # First character only (M or F)

            # Format age with gender if available
            formatted_age = patient_age
            if patient_gender and patient_age:
                formatted_age = f"{patient_age} / {patient_gender}"


            # Create header with patient info - using a table for layout
            header_table = doc.add_table(rows=2, cols=2)
            header_table.autofit = False
            
            # Name and date (first row)
            name_cell = header_table.cell(0, 0)  # Reference the first cell in the first row
            name_para = name_cell.paragraphs[0]
            name_para.space_before = Pt(0)
            name_para.space_after = Pt(0)
            name_para.paragraph_format.line_spacing = Pt(34)

            name_run = name_para.add_run(patient_name)
            name_run.bold = True
            name_run.font.size = Pt(12)
            
            date_cell = header_table.cell(0, 1)
            date_para = date_cell.paragraphs[0]
            date_para.space_after = Pt(0)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Changed from RIGHT to CENTER
            date_para.add_run(current_date)
            date_para.paragraph_format.line_spacing = Pt(34)
            
            # Address and age (second row)
            addr_cell = header_table.cell(1, 0)
            addr_para = addr_cell.paragraphs[0]
            addr_para.space_before = Pt(0) 
            addr_para.paragraph_format.left_indent = Inches(0.4)
            addr_para.add_run(patient_address)
            addr_para.paragraph_format.line_spacing = Pt(10)
            
            age_cell = header_table.cell(1, 1)
            if formatted_age:
                age_para = age_cell.paragraphs[0]
                age_para.space_before = Pt(0)
                age_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Changed from RIGHT to CENTER
                age_para.add_run(formatted_age)
                age_para.paragraph_format.line_spacing = Pt(10)
            
            
            # Add content based on print type
            if selected_type == "Prescription":
                doc.add_paragraph()  # Empty space
                
                if tree_med.get_children():
                    for idx, item in enumerate(tree_med.get_children()):
                        values = tree_med.item(item)['values']
                        brand = values[0]
                        generic = values[1]
                        quantity = values[2]
                        admin = values[3]
                        
                        # Format medication similar to PDF structure
                        generic_para = doc.add_paragraph()
                        generic_para.paragraph_format.left_indent = Inches(0.5)
                        generic_run = generic_para.add_run(f" {generic}")
                        generic_run.bold = True
                        generic_run.font.size = Pt(9)  # Set font size on run, not paragraph
                        
                        brand_para = doc.add_paragraph()
                        brand_para.paragraph_format.left_indent = Inches(0.6)
                        brand_run = brand_para.add_run(f"{brand}           #{quantity}")
                        brand_run.font.size = Pt(9)  # Set font size on run, not paragraph
                        
                        admin_para = doc.add_paragraph()
                        admin_para.paragraph_format.left_indent = Inches(0.6)
                        admin_run = admin_para.add_run(f"{admin}")
                        admin_run.font.size = Pt(9)  # Set font size on run, not paragraph
                        
                        # Add space between medications
                        if idx < len(tree_med.get_children()) - 1:
                            doc.add_paragraph()  # Add space
                else:
                    doc.add_paragraph("No medications prescribed.")
            else:  # Findings
                doc.add_paragraph()  # Empty space
                
                findings_text = text_remarks.get("1.0", tk.END).strip()
                if findings_text:
                    for line in findings_text.split('\n'):
                        if line.strip():
                            bullet_para = doc.add_paragraph()
                            bullet_para.paragraph_format.left_indent = Inches(0.2)
                            bullet_para.add_run(f"• {line}")
                else:
                    doc.add_paragraph("No findings recorded.")
            
            # Save the document
            doc.save(docx_path)
            os.startfile(docx_path)
        except Exception as e:
            messagebox.showerror("Word Error", f"Failed to create Word document: {str(e)}")

    
    # Print as Word button
    word_button = tk.Button(button_frame, text="Print as Word", bg=PRIMARY_COLOR, fg=BUTTON_TEXT_COLOR, 
                          font=("Arial", 10, "bold"), padx=20, pady=8, 
                          command=print_document_as_word)
    word_button.pack(side=tk.RIGHT, padx=5)
    
    # Regular Print button
    print_button = tk.Button(button_frame, text="Print", bg=ACCENT_COLOR, fg=BUTTON_TEXT_COLOR, 
                           font=("Arial", 10, "bold"), padx=20, pady=8, 
                           command=lambda: print_document(print_type_var.get()))
    print_button.pack(side=tk.RIGHT, padx=5)
    
    # Cancel button
    cancel_button = tk.Button(button_frame, text="Cancel", bg=WARNING_COLOR, fg=BUTTON_TEXT_COLOR,
                            font=("Arial", 10, "bold"), padx=20, pady=8, 
                            command=print_dialog.destroy)
    cancel_button.pack(side=tk.RIGHT, padx=5)
    
    # Update the preview with default selection
    update_preview()
    
    # Center the window
    print_dialog.update_idletasks()
    width = print_dialog.winfo_width()
    height = print_dialog.winfo_height()
    x = (print_dialog.winfo_screenwidth() // 2) - (width // 2)
    y = (print_dialog.winfo_screenheight() // 2) - (height // 2)
    print_dialog.geometry(f"{width}x{height}+{x}+{y}")
    
    # Make dialog modal
    print_dialog.transient(root)
    print_dialog.grab_set()

def get_formatted_prescription(patient_name):
    """Format prescription data for printing"""
    formatted_text = ""
    
    # Get current date
    current_date = datetime.now().strftime("%B %d, %Y")
    formatted_text += f"Date: {current_date}\n\n"
    
    # Get patient information
    patient_address = entry_address.get()
    patient_age = entry_age.get()
    
    # Format patient header
    formatted_text += f"{patient_name.upper()}"
    if patient_age:
        formatted_text += f"    AGE: {patient_age}"
    if patient_address:
        formatted_text += f"    ADDRESS: {patient_address}"
    formatted_text += "\n\n"
    
    # Check if there are medications
    if tree_med.get_children():
        formatted_text += "Rx:\n\n"
        
        # Add each medication to the formatted text with proper formatting
        for idx, item in enumerate(tree_med.get_children()):
            values = tree_med.item(item)['values']
            brand = values[0]
            generic = values[1]
            quantity = values[2]
            admin = values[3]
            
            # Format medication as requested
            formatted_text += f"{generic} {brand}\n"
            formatted_text += f"#{quantity} {admin}\n"
            
            # Add spacing between medications (removed separator line)
            if idx < len(tree_med.get_children()) - 1:
                formatted_text += "\n\n"
    else:
        formatted_text += "No medications prescribed."
    
    return formatted_text

def print_document(print_type):
    """Handle the actual printing process with Word document generation and direct printing"""
    try:
        print_type_var = tk.StringVar(value="Prescription")
        # Create a temporary Word document file for printing
        selected_type = print_type_var.get()
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        
        # Use python-docx to create a Word document
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Get patient information
        patient_name = entry_name.get().upper()
        patient_age = entry_age.get()
        patient_address = entry_address.get()
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Get patient gender - first character only (M or F)
        patient_gender = ""
        if gender_var.get():
            patient_gender = gender_var.get()[0]  # First character only (M or F)

        # Format age with gender if available
        formatted_age = patient_age
        if patient_gender and patient_age:
            formatted_age = f"{patient_age} / {patient_gender}"

        # Create header with patient info - using a table for layout
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Name and date (first row)
        name_cell = header_table.cell(0, 0)  # Reference the first cell in the first row
        name_para = name_cell.paragraphs[0]
        name_para.space_before = Pt(0)
        name_para.space_after = Pt(0)
        name_para.paragraph_format.line_spacing = Pt(34)

        name_run = name_para.add_run(patient_name)
        name_run.bold = True
        name_run.font.size = Pt(12)
        
        date_cell = header_table.cell(0, 1)
        date_para = date_cell.paragraphs[0]
        date_para.space_after = Pt(0)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Changed from RIGHT to CENTER
        date_para.add_run(current_date)
        date_para.paragraph_format.line_spacing = Pt(34)
        
        # Address and age (second row)
        addr_cell = header_table.cell(1, 0)
        addr_para = addr_cell.paragraphs[0]
        addr_para.space_before = Pt(0) 
        addr_para.paragraph_format.left_indent = Inches(0.4)
        addr_para.add_run(patient_address)
        addr_para.paragraph_format.line_spacing = Pt(10)
        
        age_cell = header_table.cell(1, 1)
        if formatted_age:
            age_para = age_cell.paragraphs[0]
            age_para.space_before = Pt(0)
            age_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Changed from RIGHT to CENTER
            age_para.add_run(formatted_age)
            age_para.paragraph_format.line_spacing = Pt(10)
        
        
        # Add content based on print type
        if selected_type == "Prescription":
            doc.add_paragraph()  # Empty space
             
            if tree_med.get_children():
                for idx, item in enumerate(tree_med.get_children()):
                    values = tree_med.item(item)['values']
                    brand = values[0]
                    generic = values[1]
                    quantity = values[2]
                    admin = values[3]
                    
                    # Format medication similar to PDF structure
                    generic_para = doc.add_paragraph()
                    generic_para.paragraph_format.left_indent = Inches(0.5)
                    generic_run = generic_para.add_run(f" {generic}")
                    generic_run.bold = True
                    
                    brand_para = doc.add_paragraph()
                    brand_para.paragraph_format.left_indent = Inches(0.6)
                    brand_para.add_run(f"{brand}           #{quantity}")
                    
                    admin_para = doc.add_paragraph()
                    admin_para.paragraph_format.left_indent = Inches(0.6)
                    admin_para.add_run(f"{admin}")
                    
                    # Add space between medications
                    if idx < len(tree_med.get_children()) - 1:
                        doc.add_paragraph()  # Add space
            else:
                doc.add_paragraph("No medications prescribed.")
        else:  # Findings
            doc.add_paragraph()  # Empty space
            
            findings_text = text_remarks.get("1.0", tk.END).strip()
            if findings_text:
                for line in findings_text.split('\n'):
                    if line.strip():
                        bullet_para = doc.add_paragraph()
                        bullet_para.paragraph_format.left_indent = Inches(0.2)
                        bullet_para.add_run(f"• {line}")
            else:
                doc.add_paragraph("No findings recorded.")
        
        # Save the document
        doc.save(path)
        
        # Open the Word document file for printing
        if os.name == 'nt':  # Windows
            os.startfile(path, "print")
        else:  # macOS and Linux
            subprocess.call(['lpr', path])
            
        messagebox.showinfo("Print", "Document sent to printer.")
    except Exception as e:
        messagebox.showerror("Print Error", f"Failed to print: {str(e)}")
    finally:
        # Clean up the temp file after a delay to allow printing
        def cleanup():
            try:
                os.unlink(path)
            except:
                pass
        root.after(10000, cleanup)  # 10 seconds delay
        
def open_lab_charts(new_files=None):
    """Open the Lab/Charts window with any newly selected files"""
    if not entry_name.get():
        messagebox.showwarning("Warning", "Please select a patient first.")
        return
    
    patient_name = entry_name.get()
    patient_id = patient_dict.get(patient_name)  # Get patient_id from patient_dict
    if not patient_id:
        messagebox.showerror("Error", "Patient not found in database.")
        return
    
    lab_window = LabChartsWindow(root, patient_name, patient_id, new_files)
    
def open_scan_dialog():
    file_types = [
        ('Image files', '*.png *.jpg *.jpeg *.gif *.bmp'),
        ('All files', '*.*')
    ]
    files = filedialog.askopenfilenames(
        title="Select Images",
        filetypes=file_types
    )
    
    if files:
        open_lab_charts(files)
        
# Sidebar buttons
btn_save = create_sidebar_button("Save Record", save_record, ACCENT_COLOR)
btn_update = create_sidebar_button("Update Record", update_record, PRIMARY_COLOR)
btn_clear = create_sidebar_button("Clear Form", clear_form, PRIMARY_COLOR)
btn_delete = create_sidebar_button("Delete Record", delete_record, WARNING_COLOR)
ttk.Separator(sidebar).pack(fill=tk.X, padx=10, pady=5)
btn_print = create_sidebar_button("Print", open_print_dialog, PRIMARY_COLOR)  # Add print button
btn_med_cert = create_sidebar_button("Medical Certificate", open_med_cert, PRIMARY_COLOR)
btn_scan = create_sidebar_button("Scan/Import", open_scan_dialog, PRIMARY_COLOR)
btn_lab = create_sidebar_button("Lab/Charts", open_lab_charts, PRIMARY_COLOR)

# Exit button at the bottom of sidebar
tk.Frame(sidebar, bg=SIDEBAR_COLOR, height=100).pack(fill=tk.X, expand=True)
btn_exit = create_sidebar_button("EXIT", root.quit, WARNING_COLOR)

# ----------------- Patient Information Section ----------------- #
frame_patient = tk.LabelFrame(content_frame, text="PATIENT INFORMATION", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 12, "bold"))
frame_patient.place(x=10, y=10, width=650, height=320)

# Left column of patient info
tk.Label(frame_patient, text="Name:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=0, column=0, padx=5, pady=5, sticky="w")

name_frame = tk.Frame(frame_patient, bg=SECONDARY_COLOR)
name_frame.grid(row=0, column=1, padx=5, pady=5, sticky="w")

entry_name = AutocompleteCombobox(name_frame, width=37)
entry_name.pack(side=tk.LEFT, fill=tk.X, expand=True)



tk.Label(frame_patient, text="Address:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=1, column=0, padx=5, pady=5, sticky="w")
entry_address = ttk.Entry(frame_patient, width=40)
entry_address.grid(row=1, column=1, padx=5, pady=5)

tk.Label(frame_patient, text="Birthdate:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=2, column=0, padx=5, pady=5, sticky="w")
birthdate_frame = tk.Frame(frame_patient, bg=SECONDARY_COLOR)
birthdate_frame.grid(row=2, column=1, padx=5, pady=5, sticky="w")

selected_date = tk.StringVar()
birthdate_entry = ttk.Entry(birthdate_frame, width=15, textvariable=selected_date)
birthdate_entry.pack(side=tk.LEFT, padx=(0,5))


    
def show_calendar():
    cal_window = tk.Toplevel(root)
    cal_window.title("Select Birthdate")
    cal_window.geometry("300x250")
    
    def set_date():
        date = cal.selection_get()
        selected_date.set(date.strftime('%Y-%m-%d'))
        update_age()
        cal_window.destroy()
    
    cal = Calendar(cal_window, 
                  selectmode='day',
                  date_pattern='y-mm-dd',
                  background=PRIMARY_COLOR,
                  foreground=BUTTON_TEXT_COLOR,
                  headersbackground=PRIMARY_COLOR,
                  headersforeground=BUTTON_TEXT_COLOR)
    cal.pack(padx=10, pady=10)
    
    ttk.Button(cal_window, text="Select", command=set_date).pack(pady=5)

cal_button = ttk.Button(birthdate_frame, text="📅", width=3, command=show_calendar)
cal_button.pack(side=tk.LEFT)

tk.Label(frame_patient, text="Age:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=2, column=1, padx=(200, 5), pady=5, sticky="w")
entry_age = ttk.Entry(frame_patient, width=10)
entry_age.grid(row=2, column=1, padx=(240, 5), pady=5, sticky="w")

tk.Label(frame_patient, text="Phone:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=3, column=0, padx=5, pady=5, sticky="w")
entry_phone = ttk.Entry(frame_patient, width=20)
entry_phone.grid(row=3, column=1, padx=5, pady=5, sticky="w")


frame_bp = tk.Frame(content_frame, bg=SECONDARY_COLOR)
frame_bp.place(x=10, y=335, width=485, height=40)

tk.Label(frame_bp, text="Blood Pressure:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)
entry_bp = ttk.Entry(frame_bp, width=15)
entry_bp.pack(side=tk.LEFT, padx=5)


def load_checkup_details(event=None):
    selected_date = checkup_history_var.get()
    if not selected_date or selected_date == "No previous checkups":
        return
        
    try:
        for item in tree_med.get_children():
            tree_med.delete(item)
        
        db = DatabaseHelper()
        patient_id = patient_dict[entry_name.get()]
        
        # Clear current medications in tree
        for item in tree_med.get_children():
            tree_med.delete(item)
                
        # Load prescriptions for this checkup date
        prescriptions = db.get_prescriptions_for_checkup(patient_id, selected_date)
        
        # Insert prescriptions into treeview
        for rx in prescriptions:
            tree_med.insert("", "end", values=(
                rx[0],  # brand
                rx[1],  # generic
                rx[2],  # quantity
                rx[3]   # administration
            ))
            
        # Load remarks/findings
        checkup = next((c for c in current_checkups if c[4] == selected_date), None)
        if checkup:
            text_remarks.delete("1.0", tk.END)
            text_remarks.insert("1.0", checkup[1] or "")
            
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load checkup details: {str(e)}")
        
# Add this after the Blood Pressure field in the patient information section
tk.Label(frame_patient, text="Checkup History:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=4, column=0, padx=5, pady=5, sticky="w")
checkup_history_var = tk.StringVar()
checkup_history_dropdown = ttk.Combobox(frame_patient, textvariable=checkup_history_var, width=25, state="readonly")
checkup_history_dropdown.grid(row=4, column=1, padx=5, pady=5, sticky="w")
checkup_history_dropdown.bind("<<ComboboxSelected>>", load_checkup_details)

# Right column of patient info
tk.Label(frame_patient, text="Civil Status:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=0, column=2, padx=5, pady=5, sticky="w")
status_var = tk.StringVar()
status_options = ["Single", "Married", "Divorced", "Widowed"]
status_dropdown = ttk.Combobox(frame_patient, textvariable=status_var, values=status_options, width=15)
status_dropdown.grid(row=0, column=3, padx=5, pady=5, sticky="w")

tk.Label(frame_patient, text="Gender:", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=1, column=2, padx=5, pady=5, sticky="w")
gender_var = tk.StringVar()
ttk.Radiobutton(frame_patient, text="Male", variable=gender_var, value="Male").grid(row=1, column=3, sticky="w")
ttk.Radiobutton(frame_patient, text="Female", variable=gender_var, value="Female").grid(row=1, column=3, padx=(70, 0), sticky="w")

# ----------------- Queue System ----------------- #
frame_queue = tk.LabelFrame(content_frame, text="TODAY'S QUEUE", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 12, "bold"))
frame_queue.place(x=670, y=10, width=320, height=320)

# Current date display
current_date = datetime.now()
date_string = current_date.strftime("%A, %B %d, %Y")
tk.Label(frame_queue, text=date_string, bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 10, "bold")).pack(pady=5)

tree_queue = ttk.Treeview(frame_queue, columns=("No", "Name", "Time"), show="headings", height=8)
tree_queue.heading("No", text="#")
tree_queue.heading("Name", text="Patient Name")
tree_queue.heading("Time", text="Time")
tree_queue.column("No", width=30)
tree_queue.column("Name", width=150)
tree_queue.column("Time", width=80)
tree_queue.pack(padx=5, pady=5)

# Add scrollbar to queue
scrollbar = ttk.Scrollbar(frame_queue, orient="vertical", command=tree_queue.yview)
tree_queue.configure(yscroll=scrollbar.set)
scrollbar.place(relx=0.97, relheight=0.58)

queue_frame = tk.Frame(frame_queue, bg=SECONDARY_COLOR)
queue_frame.pack(pady=5)

# Queue counter
queue_counter = 0

def add_to_queue():
    global queue_counter
    name = entry_name.get()
    if name:
        try:
            # Get current time
            current_time = datetime.now().strftime("%H:%M")
            
            # Create or verify the Queue table structure first
            db = DatabaseHelper()
            db.create_tables()
            
            # Increment queue counter
            queue_counter += 1
            
            # Save to database 
            queue_id = db.add_to_queue((queue_counter, name, current_time))
            
            if queue_id:
                # Add to treeview with the queue_id as a tag
                tree_queue.insert("", "end", values=(queue_counter, name, current_time), tags=(str(queue_id),))
                
                # Refresh patient list after adding to queue
                refresh_patient_list(name)
                
                # Show confirmation
                status_label = tk.Label(frame_queue, text="✓ Added to queue", 
                                      bg=ACCENT_COLOR, fg="white", font=("Arial", 9))
                status_label.place(relx=0.5, rely=0.92, anchor="center")
                # Remove after 2 seconds
                root.after(2000, status_label.destroy)
            else:
                # Try to recreate the Queue table and try again
                db.create_tables()
                queue_id = db.add_to_queue((queue_counter, name, current_time))
                
                if queue_id:
                    tree_queue.insert("", "end", values=(queue_counter, name, current_time), tags=(str(queue_id),))
                    refresh_patient_list(name)
                else:
                    messagebox.showerror("Database Error", "Failed to add patient to queue after multiple attempts.")
                    queue_counter -= 1  # Rollback counter increment
                
        except Exception as e:
            # Show detailed error message
            messagebox.showerror("Error", f"Could not save to queue: {str(e)}")
            queue_counter -= 1  # Rollback counter increment
    else:
        messagebox.showwarning("Input Error", "Please enter patient name first.")

def remove_from_queue():
    selected_item = tree_queue.selection()
    if selected_item:
        # Get the queue_id from the item's tags
        queue_id = tree_queue.item(selected_item, "tags")[0]
        
        try:
            db = DatabaseHelper()
            if db.remove_from_queue(queue_id):
                tree_queue.delete(selected_item)
            else:
                messagebox.showerror("Database Error", "Failed to remove patient from queue.")
        except Exception as e:
            messagebox.showerror("Error", f"Could not remove from queue: {str(e)}")
    else:
        messagebox.showwarning("Selection Error", "Select a patient to remove from queue!")

btn_add_queue = tk.Button(queue_frame, text="Add to Queue", bg=ACCENT_COLOR, fg=BUTTON_TEXT_COLOR, padx=10, pady=5, command=add_to_queue)
btn_add_queue.pack(side=tk.LEFT, padx=5)

btn_remove_queue = tk.Button(queue_frame, text="Unqueue", bg=WARNING_COLOR, fg=BUTTON_TEXT_COLOR, padx=10, pady=5, command=remove_from_queue)
btn_remove_queue.pack(side=tk.LEFT, padx=5)

# ----------------- Remarks Section ----------------- #
frame_remarks = tk.LabelFrame(content_frame, text="DIAGNOSIS", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 12, "bold"))
frame_remarks.place(x=10, y=380, width=485, height=240) 

text_remarks = tk.Text(frame_remarks, width=55, height=14)
text_remarks.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

# Add scrollbar to remarks
remarks_scrollbar = ttk.Scrollbar(frame_remarks, orient="vertical", command=text_remarks.yview)
text_remarks.configure(yscrollcommand=remarks_scrollbar.set)
remarks_scrollbar.place(relx=0.97, relheight=0.9)

# ----------------- Prescription Section ----------------- #
frame_med = tk.LabelFrame(content_frame, text="PRESCRIPTION/MEDICATION", bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 12, "bold"))
frame_med.place(x=505, y=340, width=485, height=280)

# Create treeview with updated columns
tree_med = ttk.Treeview(frame_med, columns=("Brand", "Generic", "Quantity", "Administration"), show="headings", height=8)
tree_med.heading("Brand", text="Brand Name")
tree_med.heading("Generic", text="Generic Name")
tree_med.heading("Quantity", text="Quantity")
tree_med.heading("Administration", text="Administration")
tree_med.column("Brand", width=120)
tree_med.column("Generic", width=120)
tree_med.column("Quantity", width=80)
tree_med.column("Administration", width=120)
tree_med.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

# Add scrollbar to prescription
med_scrollbar = ttk.Scrollbar(frame_med, orient="vertical", command=tree_med.yview)
tree_med.configure(yscroll=med_scrollbar.set)
med_scrollbar.place(relx=0.97, relheight=0.7)

# Button frame for medication management
med_button_frame = tk.Frame(frame_med, bg=SECONDARY_COLOR)
med_button_frame.pack(padx=1, pady=1, fill=tk.X)




def clear_prescriptions(show_confirmation=True):
    """Clear all medications from the prescription table"""
    if tree_med.get_children():  # Check if there are any prescriptions
        if not show_confirmation or messagebox.askyesno("Clear Prescriptions", "Are you sure you want to clear all prescriptions?"):
            for item in tree_med.get_children():
                tree_med.delete(item)
                
btn_clear_meds = tk.Button(med_button_frame, text="Clear All", 
                          bg=WARNING_COLOR, fg=BUTTON_TEXT_COLOR, 
                          padx=10, pady=5, command=clear_prescriptions)
btn_clear_meds.pack(side=tk.LEFT, padx=5)                

def open_medication_management():
    # Get current medications from tree_med
    current_medications = []
    for item in tree_med.get_children():
        values = tree_med.item(item)['values']
        current_medications.append({
            "id": None,  # Since we don't have ID in the prescription table
            "brand": values[0],  # Brand name
            "generic": values[1],  # Generic name
            "quantity": values[2],  # Quantity
            "administration": values[3],  # Administration
            "tree_id": item
        })

    # Define callback function to receive medications from the management window
    def update_medications(medications):
        # Clear current medications in the tree
        for item in tree_med.get_children():
            tree_med.delete(item)
        
        # Add new medications
        for med in medications:
            tree_med.insert("", "end", values=(
                med["brand"],
                med["generic"],
                med["quantity"],
                med["administration"]
            ))
    
    # Open medication management window with current medications
    med_window = MedicationManagementWindow(root, callback=update_medications)
    
    # After window is created, populate it with current medications
    for med in current_medications:
        med_window.tree.insert("", "end", values=(
            med["id"],
            med["brand"],
            med["generic"],
            med["quantity"],
            med["administration"]
        ))

def remove_selected_medication():
    selected_item = tree_med.selection()
    if selected_item:
        tree_med.delete(selected_item)
    else:
        messagebox.showwarning("Selection Error", "Select a medication to remove!")

# Add buttons for medication management
btn_manage_meds = tk.Button(med_button_frame, text="Manage Medications", 
                          bg=ACCENT_COLOR, fg=BUTTON_TEXT_COLOR, 
                          padx=10, pady=5, command=open_medication_management)
btn_manage_meds.pack(side=tk.LEFT, padx=5)

btn_remove_med = tk.Button(med_button_frame, text="Remove Selected", 
                          bg=WARNING_COLOR, fg=BUTTON_TEXT_COLOR, 
                          padx=10, pady=5, command=remove_selected_medication)
btn_remove_med.pack(side=tk.LEFT, padx=5)

# Let's update the add_to_queue function to calculate age from birthdate
def calculate_age(birth_date_str):
    try:
        birth_date = datetime.strptime(birth_date_str, "%Y-%m-%d").date()
        today = date.today()
        age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
        return str(age)
    except:
        return ""

# Add a function to update age when birthdate changes
def update_age(event=None):
    try:
        birth_date = datetime.strptime(selected_date.get(), '%Y-%m-%d')
        today = date.today()
        age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
        entry_age.delete(0, tk.END)
        entry_age.insert(0, str(age))
    except:
        pass

# Add this after the entry definitions
def load_patient_names():
    db = DatabaseHelper()
    patients = db.get_patients()
    patient_names = [patient[1] for patient in patients]
    patient_names.sort()  # Sort alphabetically
    entry_name['values'] = patient_names
    return {name: id for id, name in patients}

patient_dict = load_patient_names()

def on_name_select(event=None):
    selected_name = entry_name.get()
    if selected_name in patient_dict:
        db = DatabaseHelper()
        patient = db.get_patient_details(patient_dict[selected_name])
        load_checkup_history(patient)
        if patient:
            # Clear existing fields
            entry_address.delete(0, tk.END)
            entry_phone.delete(0, tk.END)
            entry_bp.delete(0, tk.END)
            
            # Fill in patient details from database
            entry_address.insert(0, patient[2])  # address
            # Convert string date to datetime before setting
            try:
                date_obj = datetime.strptime(patient[3], '%Y-%m-%d')
                selected_date.set(date_obj.strftime('%Y-%m-%d'))
            except:
                selected_date.set(datetime.now().strftime('%Y-%m-%d'))
                
            entry_phone.insert(0, patient[9])  # phone
            status_var.set(patient[5])  # civil_status 
            gender_var.set(patient[8])  # gender
            update_age()
            
            # Load most recent checkup info
            try:
                conn = db.get_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT blood_pressure FROM Checkups 
                    WHERE patient_id = ? 
                    ORDER BY dateOfVisit DESC LIMIT 1
                """, (patient_dict[selected_name],))
                latest_bp = cursor.fetchone()
                if (latest_bp and latest_bp[0]):
                    entry_bp.insert(0, latest_bp[0])
                conn.close()
                
                # Load checkup history for dropdown
                load_checkup_history(patient_dict[selected_name])
            except Exception as e:
                print(f"Error loading patient data: {e}")

# Modify the queue system to allow editing
def edit_queue_item(event=None):
    selected_item = tree_queue.selection()
    if selected_item:
        # Get the patient name from queue
        values = tree_queue.item(selected_item)['values']
        patient_name = values[1]
        
        # Clear current form
        clear_form(show_message=False)
        
        # Set the patient name and trigger the selection event
        entry_name.set(patient_name)
        on_name_select()  # This will load the patient details

# Bind double click on queue
tree_queue.bind('<Double-1>', edit_queue_item)

# Add this function near other birthdate-related functions
def on_date_input(event=None):
    try:
        date_str = selected_date.get()
        # Only process if there's actual input
        if date_str and date_str != "YYYY-MM-DD":
            # Try to parse the date
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            # If successful, update age
            today = date.today()
            age = today.year - date_obj.year - ((today.month, today.day) < (date_obj.month, date_obj.day))
            
            # Safely update the age entry
            entry_age.delete(0, 'end')  # Use string 'end' instead of tk.END
            entry_age.insert(0, str(age))  # Insert at position 0
            
            birthdate_entry.config(foreground="black")
    except ValueError:
        # If date is invalid, show in red but don't update age
        birthdate_entry.config(foreground="red")

# Near where birthdate_entry is created, add these bindings
birthdate_entry.bind("<KeyRelease>", on_date_input)
birthdate_entry.bind("<FocusOut>", on_date_input)

# Add placeholder text and handlers
birthdate_entry.insert(0, "YYYY-MM-DD")
birthdate_entry.config(foreground="gray")

# Fix the placeholder text handlers
def on_entry_focus_in(event):
    if birthdate_entry.get() == "YYYY-MM-DD":
        birthdate_entry.delete(0, 'end')  # Use string 'end' instead of tk.END
        birthdate_entry.config(foreground="black")

def on_entry_focus_out(event):
    if not birthdate_entry.get():
        birthdate_entry.delete(0, 'end')  # Use string 'end' instead of tk.END
        birthdate_entry.insert(0, "YYYY-MM-DD")
        birthdate_entry.config(foreground="gray")
    else:
        on_date_input(event)

birthdate_entry.bind("<FocusIn>", on_entry_focus_in)
birthdate_entry.bind("<FocusOut>", on_entry_focus_out)

# Add this function to refresh the patient list
def refresh_patient_list(keep_selection=None):
    """
    Refresh the patient dropdown list with current data from the database.
    If keep_selection is provided, will maintain that patient as selected.
    """
    global patient_dict
    # Store current selection if needed
    current_selection = None
    if keep_selection:
        current_selection = keep_selection
    elif entry_name.get():
        current_selection = entry_name.get()
    
    # Reload patient names from database
    patient_dict = load_patient_names()
    
    # Restore the selection if needed
    if current_selection:
        entry_name.set(current_selection)

# Call refresh_patient_list after the window is loaded
root.after(100, refresh_patient_list)

# Add these new functions to handle checkup history
def load_checkup_history(patient_id):
    """Load the checkup history for a patient into the dropdown"""
    try:
        db = DatabaseHelper()
        checkups = db.get_patient_checkups(patient_id)
        
        # Format dates for dropdown
        checkup_dates = [checkup[4] for checkup in checkups]  # Using last_checkup_date
        
        # Update dropdown values
        if checkup_dates:
            checkup_history_dropdown['values'] = checkup_dates
            checkup_history_dropdown.set(checkup_dates[0])  # Set to most recent
            global current_checkups
            current_checkups = checkups
        else:
            checkup_history_dropdown['values'] = ["No previous checkups"]
            checkup_history_dropdown.set("No previous checkups")
            
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load checkup history: {str(e)}")


def show_checkup_notification(checkup):
    """Display checkup details in a notification window"""
    notification = tk.Toplevel(root)
    notification.title("Checkup History Details")
    notification.geometry("600x500")  # Larger window size
    notification.configure(bg=SECONDARY_COLOR)
    
    # Get checkup details
    checkup_id = checkup[0]
    findings = checkup[1] or "No findings recorded"
    lab_ids = checkup[2] or "No labs recorded"
    checkup_date = checkup[3]
    last_checkup_date = checkup[4]
    blood_pressure = checkup[5] or "Not recorded"
    
    # Create header frame
    header_frame = tk.Frame(notification, bg=PRIMARY_COLOR, padx=15, pady=10)
    header_frame.pack(fill=tk.X)
    
    # Add header content
    tk.Label(header_frame, text=f"Checkup Details - {checkup_date}", 
             bg=PRIMARY_COLOR, fg="white", font=("Arial", 14, "bold")).pack(anchor="w")
    
    # Create main content frame with padding
    main_frame = tk.Frame(notification, bg=SECONDARY_COLOR, padx=20, pady=15)
    main_frame.pack(fill=tk.BOTH, expand=True)
    
    # Patient details section if available
    if entry_name.get() in patient_dict:
        patient_id = patient_dict[entry_name.get()]
        db = DatabaseHelper()
        patient = db.get_patient_details(patient_id)
        if patient:
            patient_frame = tk.LabelFrame(main_frame, text="Patient Information", 
                                       bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 11, "bold"),
                                       padx=10, pady=10)
            patient_frame.pack(fill=tk.X, pady=(0, 15))
            
            info_grid = tk.Frame(patient_frame, bg=SECONDARY_COLOR)
            info_grid.pack(fill=tk.X)
            
            # Row 1
            tk.Label(info_grid, text="Name:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, 
                   font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", padx=(0, 5), pady=2)
            tk.Label(info_grid, text=patient[1], bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=0, column=1, sticky="w", padx=5, pady=2)
            
            tk.Label(info_grid, text="Gender:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, 
                   font=("Arial", 10, "bold")).grid(row=0, column=2, sticky="w", padx=(20, 5), pady=2)
            tk.Label(info_grid, text=patient[8], bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=0, column=3, sticky="w", padx=5, pady=2)
            
            # Row 2
            tk.Label(info_grid, text="Age at checkup:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, 
                   font=("Arial", 10, "bold")).grid(row=1, column=0, sticky="w", padx=(0, 5), pady=2)
            
            # Calculate age at time of checkup
            try:
                birth_date = datetime.strptime(patient[3], '%Y-%m-%d').date()
                checkup_datetime = datetime.strptime(checkup_date, '%Y-%m-%d').date()
                age_at_checkup = checkup_datetime.year - birth_date.year - ((checkup_datetime.month, checkup_datetime.day) < (birth_date.month, birth_date.day))
                tk.Label(info_grid, text=str(age_at_checkup), bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=1, column=1, sticky="w", padx=5, pady=2)
            except:
                tk.Label(info_grid, text="N/A", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=1, column=1, sticky="w", padx=5, pady=2)
            
            tk.Label(info_grid, text="Status:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, 
                   font=("Arial", 10, "bold")).grid(row=1, column=2, sticky="w", padx=(20, 5), pady=2)
            tk.Label(info_grid, text=patient[5] or "Not recorded", bg=SECONDARY_COLOR, fg=TEXT_COLOR).grid(row=1, column=3, sticky="w", padx=5, pady=2)
    
    # Vital signs section
    vitals_frame = tk.LabelFrame(main_frame, text="Vital Signs", 
                                bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 11, "bold"),
                                padx=10, pady=10)
    vitals_frame.pack(fill=tk.X, pady=(0, 15))
    
    vitals_grid = tk.Frame(vitals_frame, bg=SECONDARY_COLOR)
    vitals_grid.pack(fill=tk.X)
    
    tk.Label(vitals_grid, text="Blood Pressure:", bg=SECONDARY_COLOR, fg=TEXT_COLOR, 
           font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", padx=(0, 5), pady=5)
    tk.Label(vitals_grid, text=blood_pressure, bg=SECONDARY_COLOR, fg=TEXT_COLOR,
           font=("Arial", 10)).grid(row=0, column=1, sticky="w", padx=5, pady=5)
    
    # Findings & Remarks with scrollable text
    findings_frame = tk.LabelFrame(main_frame, text="Findings & Remarks", 
                                 bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 11, "bold"),
                                 padx=10, pady=10)
    findings_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
    
    # Create a frame to hold the text widget and scrollbar
    text_container = tk.Frame(findings_frame, bg=SECONDARY_COLOR)
    text_container.pack(fill=tk.BOTH, expand=True)
    
    # Text widget for findings with proper styling
    findings_text = tk.Text(text_container, wrap=tk.WORD, width=60, height=8,
                          font=("Arial", 10), padx=8, pady=8,
                          bg="white", fg=TEXT_COLOR)
    findings_text.insert("1.0", findings)
    findings_text.config(state="disabled")  # Make read-only
    
    # Add scrollbar to findings text
    text_scroll = ttk.Scrollbar(text_container, orient="vertical", command=findings_text.yview)
    findings_text.configure(yscrollcommand=text_scroll.set)
    
    findings_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    text_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    
    # Try to get prescriptions for this checkup
    try:
        db = DatabaseHelper()
        conn = db.get_connection()
        cursor = conn.cursor()
        
        # Check if there are prescriptions with the same date
        cursor.execute("""
            SELECT brand, generic, quantity, administration
            FROM Prescriptions 
            WHERE patient_id = (
                SELECT patient_id FROM Checkups WHERE id = ?
            ) AND last_checkup_date = ?
        """, (checkup_id, checkup_date))
        
        prescriptions = cursor.fetchall()
        conn.close()
        
        # Add prescriptions section if there are any
        if prescriptions:
            rx_frame = tk.LabelFrame(main_frame, text="Prescribed Medications", 
                                   bg=SECONDARY_COLOR, fg=TEXT_COLOR, font=("Arial", 11, "bold"),
                                   padx=10, pady=10)
            rx_frame.pack(fill=tk.X, pady=(0, 15))
            
            # Create a container for the treeview and scrollbar
            rx_container = tk.Frame(rx_frame, bg=SECONDARY_COLOR)
            rx_container.pack(fill=tk.X, expand=True)
            
            # Create a treeview for prescriptions
            rx_tree = ttk.Treeview(rx_container, columns=("Brand", "Generic", "Quantity", "Administration"), 
                                  show="headings", height=min(len(prescriptions), 3))
            rx_tree.heading("Brand", text="Brand Name")
            rx_tree.heading("Generic", text="Generic Name")
            rx_tree.heading("Quantity", text="Qty")
            rx_tree.heading("Administration", text="Administration")
            
            rx_tree.column("Brand", width=150)
            rx_tree.column("Generic", width=150)
            rx_tree.column("Quantity", width=50)
            rx_tree.column("Administration", width=150)
            
            # Add prescriptions to the tree
            for rx in prescriptions:
                rx_tree.insert("", tk.END, values=rx)
            
            # Add scrollbar to the treeview
            rx_scroll = ttk.Scrollbar(rx_container, orient="vertical", command=rx_tree.yview)
            rx_tree.configure(yscrollcommand=rx_scroll.set)
            
            rx_tree.pack(side=tk.LEFT, fill=tk.X, expand=True)
            rx_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    except Exception as e:
        print(f"Error loading prescriptions: {e}")
    
    # Button frame
    button_frame = tk.Frame(notification, bg=SECONDARY_COLOR, padx=15, pady=15)
    button_frame.pack(fill=tk.X)
    
    # Close button
    close_button = tk.Button(button_frame, text="Close", bg=PRIMARY_COLOR, fg=BUTTON_TEXT_COLOR,
                           font=("Arial", 10, "bold"), padx=20, pady=8, command=notification.destroy)
    close_button.pack(side=tk.RIGHT, padx=10)
    
    # Copy to Remarks button
    def copy_to_remarks():
        text_remarks.delete("1.0", tk.END)
        text_remarks.insert("1.0", findings)
        # Show confirmation toast
        confirmation = tk.Label(notification, text="✓ Copied to remarks", bg="green", fg="white", 
                              font=("Arial", 10, "bold"), padx=10, pady=5)
        confirmation.place(relx=0.5, rely=0.9, anchor="center")
        notification.after(2000, confirmation.destroy)  # Remove after 2 seconds
    
    copy_button = tk.Button(button_frame, text="Use as Template", bg=ACCENT_COLOR, fg=BUTTON_TEXT_COLOR,
                          font=("Arial", 10, "bold"), padx=20, pady=8, command=copy_to_remarks)
    copy_button.pack(side=tk.RIGHT, padx=10)
    
    # Apply blood pressure button
    def apply_bp():
        entry_bp.delete(0, tk.END)
        if blood_pressure and blood_pressure != "Not recorded":
            entry_bp.insert(0, blood_pressure)
        # Show confirmation toast
        confirmation = tk.Label(notification, text="✓ Blood pressure applied", bg="green", fg="white", 
                              font=("Arial", 10, "bold"), padx=10, pady=5)
        confirmation.place(relx=0.5, rely=0.9, anchor="center")
        notification.after(2000, confirmation.destroy)  # Remove after 2 seconds
    
    bp_button = tk.Button(button_frame, text="Apply BP", bg=PRIMARY_COLOR, fg=BUTTON_TEXT_COLOR,
                        font=("Arial", 10, "bold"), padx=20, pady=8, command=apply_bp)
    bp_button.pack(side=tk.RIGHT, padx=10)
    
    # Center the window on the screen
    notification.update_idletasks()
    width = notification.winfo_width()
    height = notification.winfo_height()
    x = (notification.winfo_screenwidth() // 2) - (width // 2)
    y = (notification.winfo_screenheight() // 2) - (height // 2)
    notification.geometry(f"{width}x{height}+{x}+{y}")
    
    # Make window modal
    notification.transient(root)
    notification.grab_set()
    root.wait_window(notification)

# Add a global variable to store current checkups data
current_checkups = []

# Initialize queue from database
try:
    # Initialize the database first
    db = DatabaseHelper()
    db.create_tables()
    
    # Load today's queue
    load_today_queue()
    
    # Clean up old queue entries (older than 7 days)
    db.clear_old_queue(7)
except Exception as e:
    print(f"Error initializing queue: {e}")
    
root.mainloop()
