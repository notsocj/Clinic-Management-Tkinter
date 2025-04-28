import tkinter as tk
from tkinter import ttk, font, messagebox, filedialog
from datetime import datetime
import tkinter.scrolledtext as scrolledtext
import os
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
import subprocess

class MedicalCertificateWindow:
    def __init__(self, parent, patient_data=None):
        self.window = tk.Toplevel(parent)
        self.window.title("Medical Certificate")
        self.window.geometry("800x600")
        self.window.configure(bg="#f0f0f0")
        
        # Parent reference
        self.parent = parent
        
        # Patient data should include name, age, address, findings, remarks
        self.patient_data = patient_data or {
            "name": "", 
            "age": "", 
            "address": "", 
            "findings": "",
            "remarks": ""
        }
        
        self.create_widgets()
        self.generate_medical_certificate()
        
    def create_widgets(self):
        # Main frame
        main_frame = tk.Frame(self.window, bg="#f0f0f0")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Instructions label
        instructions = tk.Label(main_frame, text="You can edit the certificate text below before printing or saving",
                              bg="#f0f0f0", fg="#3498db", font=("Arial", 10, "italic"))
        instructions.pack(fill=tk.X, pady=(0, 5))
        
        # Rich text box with scrollbar - now editable by default
        self.rich_text = scrolledtext.ScrolledText(main_frame, wrap=tk.WORD, width=80, height=30)
        self.rich_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Add proper Enter key binding to ensure newlines work correctly
        self.rich_text.bind("<Return>", self.handle_return_key)
        
        # Button frame
        button_frame = tk.Frame(self.window, bg="#f0f0f0")
        button_frame.pack(fill=tk.X, padx=20, pady=(0, 10))
        
        # Save to File button (new)
        self.save_button = ttk.Button(button_frame, text="Save to File", command=self.save_to_file)
        self.save_button.pack(side=tk.LEFT, padx=5)
        
        # Print button (renamed from Print Certificate to Export as PDF)
        self.print_button = ttk.Button(button_frame, text="Export as PDF", command=self.export_as_pdf)
        self.print_button.pack(side=tk.LEFT, padx=5)
        
        # PDF Print button - new button for generating and printing PDF
        self.pdf_button = ttk.Button(button_frame, text="Print as PDF", command=self.print_as_pdf)
        self.pdf_button.pack(side=tk.LEFT, padx=5)
        
        # Word Print button - new button for printing as Word
        self.word_button = ttk.Button(button_frame, text="Print as Word", command=self.print_as_word)
        self.word_button.pack(side=tk.LEFT, padx=5)
        
        # Exit button
        self.exit_button = ttk.Button(button_frame, text="Close", command=self.window.destroy)
        self.exit_button.pack(side=tk.RIGHT, padx=5)
    
    def handle_return_key(self, event=None):
        """Handle Enter key press to ensure proper newline insertion"""
        # Get current cursor position
        current_position = self.rich_text.index(tk.INSERT)
        
        # Let the normal Return key behavior happen (don't prevent default)
        return None  # This allows the default Enter behavior to work
    
    def save_to_file(self):
        """Save the certificate text to a file"""
        try:
            # Get content from the rich text widget
            content = self.rich_text.get("1.0", tk.END)
            
            # Get the patient name for the default filename
            patient_name = self.patient_data.get('name', 'Unknown').replace(' ', '_')
            default_filename = f"Medical_Certificate_{patient_name}_{datetime.now().strftime('%Y%m%d')}.txt"
            
            # Open file dialog to select save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
                initialfile=default_filename
            )
            
            if file_path:
                with open(file_path, 'w') as file:
                    file.write(content)
                messagebox.showinfo("Success", f"Certificate saved to:\n{file_path}")
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save file: {str(e)}")
    
    def export_as_pdf(self):
        """Export the certificate as a PDF file"""
        try:
            # Get the patient name for the default filename
            patient_name = self.patient_data.get('name', 'Unknown').replace(' ', '_')
            default_filename = f"Medical_Certificate_{patient_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
            
            # Open file dialog to select save location
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                initialfile=default_filename
            )
            
            if file_path:
                # Create a PDF document with A4 size
                self.generate_pdf(file_path)
                
                # Show success message with option to open the file
                response = messagebox.askyesno(
                    "Success", 
                    f"Certificate saved as PDF to:\n{file_path}\n\nWould you like to open it now?",
                    icon='info'
                )
                
                if response:
                    if os.name == 'nt':  # Windows
                        os.startfile(file_path)
                    else:  # macOS and Linux
                        subprocess.call(['xdg-open', file_path])
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create PDF: {str(e)}")
    
    def print_certificate(self):
        """Deprecated: This function is kept for backward compatibility"""
        self.export_as_pdf()
    
    def print_as_pdf(self):
        """Generate a PDF and open it for printing"""
        try:
            # Create a temporary PDF file
            fd, pdf_path = tempfile.mkstemp(suffix='.pdf')
            os.close(fd)
            
            # Generate the PDF
            self.generate_pdf(pdf_path)
            
            # Show a success message
            messagebox.showinfo("PDF Created", 
                "The medical certificate has been created as a PDF.\n"
                "It will open in your default PDF viewer where you can print it.")
            
            # Open with system default PDF viewer
            if os.name == 'nt':  # Windows
                os.startfile(pdf_path)
            else:  # macOS and Linux
                subprocess.call(['xdg-open', pdf_path])
                
            # Clean up temp file after delay
            self.window.after(30000, lambda: self.cleanup_temp_file(pdf_path))
                
        except Exception as e:
            messagebox.showerror("PDF Error", f"Failed to create PDF: {str(e)}")
    
    def print_as_word(self):
        """Generate a Word document of the medical certificate and print it"""
        try:
            # Create a temporary Word document file for printing
            fd, path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            
            # Use python-docx to create a Word document
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Get patient information from self.patient_data
            patient_name = self.patient_data.get('name', '').upper()
            patient_age = self.patient_data.get('age', '')
            patient_address = self.patient_data.get('address', '')
            findings = self.patient_data.get('findings', '')
            current_date = datetime.now().strftime("%A, %d %B %Y")
            
            # Set document margins (similar to PDF)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.1)
                section.bottom_margin = Inches(0.1)
                section.left_margin = Inches(0.1)
                section.right_margin = Inches(0.1)
            
            # 1. Header - Doctor name in blue
            header_para = doc.add_paragraph()
            header_run = header_para.add_run("DR. BELINDA O. CABAUATAN")
            header_run.bold = True
            header_run.font.size = Pt(18)
            header_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue color
            
            # 2. Add specialization - using tight spacing
            spec_para1 = doc.add_paragraph()
            spec_para1.paragraph_format.space_before = Pt(0)
            spec_para1.paragraph_format.space_after = Pt(0)
            spec_para1.add_run("ADULT NEUROLOGY").bold = True
            
            spec_para2 = doc.add_paragraph()
            spec_para2.paragraph_format.space_before = Pt(0)
            spec_para2.paragraph_format.space_after = Pt(0)
            spec_para2.add_run("BRAIN, SPINAL CORD, NERVE AND MUSCLE SPECIALIST").bold = True
            
            # 3. Medical Certificate heading
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(6)
            title_para.paragraph_format.left_indent = Inches(1.5)
            title_run = title_para.add_run("MEDICAL CERTIFICATE")
            title_run.bold = True
            title_run.font.size = Pt(18)
            
            # 4. Patient info
            doc.add_paragraph()  # Small space
            
            patient_para1 = doc.add_paragraph()
            patient_para1.add_run(f"This is to certify that {patient_name} {patient_age} years old, from {patient_address} was.")
            
            patient_para2 = doc.add_paragraph()
            patient_para2.add_run(f"seen and examined in this clinic on {current_date}.")
            
            # 5. Diagnosis section
            doc.add_paragraph()  # Space before diagnosis
            diag_title = doc.add_paragraph()
            diag_title.add_run("DIAGNOSIS:").bold = True
            diag_title.runs[0].font.size = Pt(12)
            
            # Get diagnosis text from rich text editor instead of patient data
            # to include any edits made by the user
            content_lines = self.rich_text.get("1.0", tk.END).splitlines()
            
            # Find diagnosis section
            diagnosis_text = ""
            in_diagnosis = False
            for line in content_lines:
                if "DIAGNOSIS:" in line:
                    in_diagnosis = True
                    continue
                if in_diagnosis and "This certification is issued" in line:
                    in_diagnosis = False
                    break
                if in_diagnosis and line.strip():
                    diagnosis_text += line + "\n"
            
            # Add diagnosis content
            if diagnosis_text:
                diag_para = doc.add_paragraph()
                diag_para.add_run(diagnosis_text)
            else:
                # If no diagnosis text found, use the patient data
                diag_para = doc.add_paragraph()
                diag_para.add_run(findings)
            
            # 6. Add spacing
            for _ in range(3):
                doc.add_paragraph()
            
            # 7. Certification notice
            cert_para1 = doc.add_paragraph()
            cert_para1.add_run("This certification is issued for reference use only.")
            
            cert_para2 = doc.add_paragraph()
            cert_para2.add_run(f"Issued this {current_date} at Iriga Clinic.")
            
            # 8. Add spacing before signature
            for _ in range(2):
                doc.add_paragraph()
            
            # 9. Doctor signature - right aligned
            sign_para1 = doc.add_paragraph()
            sign_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sign_para1.add_run("DR. BELINDA O. CABAUATAN M.D.")
            
            sign_para2 = doc.add_paragraph()
            sign_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sign_para2.add_run("NEUROLOGY")
            
            sign_para3 = doc.add_paragraph()
            sign_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sign_para3.add_run("Lic. No.    109769")
            
            sign_para4 = doc.add_paragraph()
            sign_para4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sign_para4.add_run("PTR. No.    4813787")
            
            # Save the document
            doc.save(path)
            
            # Show a success message
            messagebox.showinfo("Word Document Created", 
                "The medical certificate has been created as a Word document.\n"
                "It will open for printing.")
            
            # Open the Word document file for printing
            if os.name == 'nt':  # Windows
                os.startfile(path, "print")
            else:  # macOS and Linux
                subprocess.call(['lpr', path])
            
            # Clean up temp file after delay
            self.window.after(10000, lambda: self.cleanup_temp_file(path))
                
        except Exception as e:
            messagebox.showerror("Word Error", f"Failed to create Word document: {str(e)}")
    
    def cleanup_temp_file(self, file_path):
        """Clean up temporary files after a delay"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error cleaning up temporary file: {e}")
    
    def generate_pdf(self, pdf_path):
        """Generate a PDF file with the certificate content that matches the visual layout"""
        # Create a PDF document with A4 size
        edited_content = self.rich_text.get("1.0", tk.END)
        
        # Parse the edited content to extract sections
        # This is a simplistic approach - you might need more sophisticated parsing
        content_lines = edited_content.splitlines()
        
        # Create a PDF document with A4 size
        doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
        
        # Styles for different text parts
        styles = getSampleStyleSheet()
        
        # Custom styles that match the rich text display
        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=styles['Heading1'],
            fontSize=18,
            fontName='Helvetica-Bold',
            textColor=colors.blue,
            spaceAfter=0,
            leading=18     # Reduced from 20 to 18
        )
            
        subheader_style = ParagraphStyle(
            'SubheaderStyle',
            parent=styles['Heading2'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceAfter=0,
            leading=11     # Reduced from 13 to 11
        )
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=18,
            fontName='Helvetica-Bold',
            alignment=0,  # Left alignment to match rich text
            spaceAfter=6,  # Reduced from 10 to 6
            spaceBefore=3, # Reduced from 6 to 3
            leading=20     # Reduced from 22 to 20
        )
        
        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=2,  # Reduced from 4 to 2
            allowWidows=0,
            allowOrphans=0,
            leading=12     # Added lower leading value
        )
        
        diagnosis_style = ParagraphStyle(
            'DiagnosisStyle',
            parent=styles['Heading3'],
            fontSize=12,
            fontName='Helvetica',
            spaceAfter=3,  # Reduced from 6 to 3
            spaceBefore=3, # Reduced from 6 to 3
            leading=14     # Added leading value
        )
        
        signature_style = ParagraphStyle(
            'SignatureStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=2  # Right alignment
        )
        
        # Get today's date for the certificate
        today_date = datetime.now().strftime("%A, %d %B %Y")
        
        # Elements to add to the PDF
        elements = []
        
        header_index = 0
        title_index = None
        diagnosis_index = None
        signature_index = None
        
        # Find key sections by looking for marker text
        for i, line in enumerate(content_lines):
            if "MEDICAL CERTIFICATE" in line:
                title_index = i
            elif "DIAGNOSIS:" in line:
                diagnosis_index = i
            elif "DR. BELINDA O. CABAUATAN M.D." in line:
                signature_index = i
        
        # Add a new section finder for "REMARKS:"
        remarks_index = None
        for i, line in enumerate(content_lines):
            if "REMARKS:" in line:
                remarks_index = i
                break
        
        # Add a new section finder for certification text
        certification_start_idx = None
        for i, line in enumerate(content_lines):
            if "This certification is issued" in line:
                certification_start_idx = i
                break
    
        # Header section with reduced vertical spacing
        elements.append(Paragraph("DR. BELINDA O. CABAUATAN", header_style))
        elements.append(Paragraph("ADULT NEUROLOGY", subheader_style))
        elements.append(Paragraph("BRAIN, SPINAL CORD, NERVE AND MUSCLE SPECIALIST", subheader_style))
        
        # Further reduced space after the header section
        elements.append(Spacer(1, 5))  # Reduced from 8 to 5
        
        # Title - keep the same indentation as in the rich text
        elements.append(Paragraph("               MEDICAL CERTIFICATE", title_style))
        elements.append(Spacer(1, 12))
        
        # Patient information (extract from edited content)
        patient_info = ""
        if title_index is not None and diagnosis_index is not None:
            for i in range(title_index + 1, diagnosis_index):
                if content_lines[i].strip():
                    patient_info += content_lines[i].strip() + " "
        
        if patient_info:
            # Split patient info into paragraphs
            patient_info_parts = patient_info.split("seen and examined")
            if len(patient_info_parts) > 0:
                elements.append(Paragraph(patient_info_parts[0], normal_style))
            if len(patient_info_parts[1:]):
                elements.append(Paragraph("seen and examined" + patient_info_parts[1], normal_style))
        
        # Further reduce spacing between patient information and diagnosis section
        elements.append(Spacer(1, 5))  # Reduced from 8 to 5
        
        # Diagnosis section
        elements.append(Paragraph("DIAGNOSIS:", diagnosis_style))
        diagnosis_content = ""
        
        # Find where the diagnosis ends - either at REMARKS or certification 
        diagnosis_end_idx = None
        if remarks_index is not None:
            diagnosis_end_idx = remarks_index
        elif certification_start_idx is not None:
            diagnosis_end_idx = certification_start_idx
        
        # Extract diagnosis content
        if diagnosis_index is not None and diagnosis_end_idx is not None:
            # Process each line and maintain empty lines for spacing
            for i in range(diagnosis_index + 1, diagnosis_end_idx):
                current_line = content_lines[i]
                # Convert each line to HTML paragraph for proper line breaks
                if current_line.strip():
                    # Add non-empty line
                    diagnosis_content += current_line + "<br/>"
                else:
                    # Add empty line (preserves spacing)
                    diagnosis_content += "<br/>"

        if diagnosis_content.strip():
            # Create paragraph with HTML-formatted content
            elements.append(Paragraph(diagnosis_content, normal_style))
        else:
            elements.append(Spacer(1, 40))

        # Add adjusted space since we removed the remarks section
        elements.append(Spacer(1, 40))  # Reduced from 60 to 40
        
        # Find the certification lines
        certification_text = []
        if certification_start_idx is not None and signature_index is not None:
            for i in range(certification_start_idx, signature_index):
                if content_lines[i].strip():
                    certification_text.append(content_lines[i])
        
        for line in certification_text:
            if line.strip():
                elements.append(Paragraph(line, normal_style))

        # Add space before signature (reduced)
        elements.append(Spacer(1, 50))  # Reduced from 72 to 50
        
        # Doctor's signature section from the edited content
        if signature_index is not None:
            for i in range(signature_index, min(signature_index + 4, len(content_lines))):
                elements.append(Paragraph(content_lines[i].strip(), signature_style))
        
        # Build the PDF
        doc.build(elements)
        
        return pdf_path
    
    def append_text(self, text, style="normal", alignment="left", size=10, color=None, font_name="Arial"):
        # Create a unique tag name for this text style
        tag_name = f"{style}_{alignment}_{size}_{color}_{font_name}"
        
        # Configure the font
        if style == "bold":
            text_font = font.Font(family=font_name, size=size, weight="bold")
        else:
            text_font = font.Font(family=font_name, size=size)
        
        # Configure the tag with the font and alignment
        self.rich_text.tag_configure(tag_name, font=text_font)
        
        # Set alignment
        if alignment == "left":
            self.rich_text.tag_configure(tag_name, justify="left")
        elif alignment == "right":
            self.rich_text.tag_configure(tag_name, justify="right")
        elif alignment == "center":
            self.rich_text.tag_configure(tag_name, justify="center")
        
        # Set text color if specified
        if color:
            self.rich_text.tag_configure(tag_name, foreground=color)
        
        # Handle multi-line text correctly - split by newlines and insert separately
        lines = text.split("\n")
        for i, line in enumerate(lines):
            self.rich_text.insert(tk.END, line, tag_name)
            # Add a newline after each line except the last one if there are multiple lines
            if i < len(lines) - 1 or text.endswith("\n"):
                self.rich_text.insert(tk.END, "\n")
        
        # Always add a final newline after the entire text
        self.rich_text.insert(tk.END, "\n")
    
    def generate_medical_certificate(self):
        # Clear previous content
        self.rich_text.delete(1.0, tk.END)
        
        # Get today's date
        today_date = datetime.now().strftime("%A, %d %B %Y")
        
        # Append the doctor's name in blue
        self.append_text("DR. BELINDA O. CABAUATAN", style="bold", alignment="left", size=18, color="blue")
        
        # Append specialization
        self.append_text("ADULT NEUROLOGY", style="bold", alignment="left", size=11)
        self.append_text("BRAIN, SPINAL CORD, NERVE AND MUSCLE SPECIALIST", style="bold", alignment="left", size=11)
        
        # Empty line
        self.rich_text.insert(tk.END, "\n")
        
        # Append the heading with Algerian font (using Arial as fallback if Algerian not available)
        self.append_text("               MEDICAL CERTIFICATE", style="bold", alignment="left", size=18, font_name="Algerian")
        
        # Empty line
        self.rich_text.insert(tk.END, "\n")
        
        # Patient information
        self.append_text(f"This is to certify that {self.patient_data['name']} {self.patient_data['age']} years old, from {self.patient_data['address']} was.", 
                        style="normal", alignment="left", size=10)
        self.append_text(f"seen and examined in this clinic on {today_date}.", 
                        style="normal", alignment="left", size=10)
        
        # Empty line
        self.rich_text.insert(tk.END, "\n")
        
        # Diagnosis
        self.append_text("DIAGNOSIS:", style="normal", alignment="left", size=12)
        self.append_text(self.patient_data.get('findings', ''), style="normal", alignment="left", size=10)
        
        # Add space - more space since we're removing remarks section
        for _ in range(12):
            self.rich_text.insert(tk.END, "\n")
        
        # Certification notice
        self.append_text("This certification is issued for reference use only.", 
                        style="normal", alignment="left", size=10)
        self.append_text(f"Issued this {today_date} at Iriga Clinic.", 
                        style="normal", alignment="left", size=10)
        
        # Add space
        for _ in range(8):
            self.rich_text.insert(tk.END, "\n")
        
        # Doctor info with right alignment
        self.append_text("DR. BELINDA O. CABAUATAN M.D.", style="normal", alignment="right", size=10)
        self.append_text("NEUROLOGY", style="normal", alignment="right", size=10)
        self.append_text("Lic. No.    109769", style="normal", alignment="right", size=10)
        self.append_text("PTR. No.    4813787", style="normal", alignment="right", size=10)
